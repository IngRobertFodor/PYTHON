"""Testy pre PDF Password Converter"""
import os, sys, io, shutil
from datetime import datetime

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
_dir = os.path.dirname(os.path.abspath(__file__))
if _dir not in sys.path:
    sys.path.insert(0, _dir)

from docx import Document
from pypdf import PdfReader
from converter import (
    get_employees, add_employee, remove_employee,
    get_employee_password, set_employee_password,
    generate_output_filename, convert_file, docx_to_pdf,
    get_output_folder, get_app_dir, MESIACE, KEYRING_SERVICE
)
import keyring

TS, TP = "TS", "test123"

def setup():
    add_employee(TS, TP)

def cleanup():
    remove_employee(TS)

# === ENV TESTY ===

def test_env_imports():
    """Vsetky importy funguju."""
    print("\n[TEST 1] Importy")
    import tkinter, pypdf, docx, reportlab
    assert keyring is not None
    print("  [OK] PASSED")

def test_env_fonts():
    """Arial font existuje."""
    print("\n[TEST 2] Arial font")
    fonts_dir = os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts')
    assert os.path.exists(os.path.join(fonts_dir, 'arial.ttf'))
    print("  [OK] PASSED")

def test_env_output_folder():
    """Vystupny priecinok sa vytvori."""
    print("\n[TEST 3] Vystupny priecinok")
    out = get_output_folder()
    if os.path.exists(out):
        shutil.rmtree(out)
    out2 = get_output_folder()
    assert os.path.exists(out2) and os.path.isdir(out2)
    print("  [OK] PASSED")

def test_env_keyring():
    """Keyring funguje (zapis/citanie)."""
    print("\n[TEST 4] Keyring R/W")
    keyring.set_password("test_svc", "test_key", "test_val")
    assert keyring.get_password("test_svc", "test_key") == "test_val"
    keyring.delete_password("test_svc", "test_key")
    print("  [OK] PASSED")

# === ZAMESTNANEC TESTY ===

def test_emp_crud():
    """Pridanie, heslo, zmena, overenie."""
    print("\n[TEST 5] Zamestnanec CRUD")
    assert TS in get_employees()
    assert get_employee_password(TS) == TP
    set_employee_password(TS, "new")
    assert get_employee_password(TS) == "new"
    set_employee_password(TS, TP)
    print("  [OK] PASSED")

def test_emp_multiple():
    """Viacero zamestnancov."""
    print("\n[TEST 6] Viacero zamestnancov")
    add_employee("AB", "a1")
    add_employee("CD", "c2")
    emps = get_employees()
    assert "AB" in emps and "CD" in emps
    assert get_employee_password("AB") == "a1"
    remove_employee("AB")
    remove_employee("CD")
    assert "AB" not in get_employees()
    print("  [OK] PASSED")

# === NAZOV SUBORU TESTY ===

def test_name_auto():
    """Automaticky nazov (predchadzajuci mesiac)."""
    print("\n[TEST 7] Nazov - auto")
    name = generate_output_filename(TS)
    now = datetime.now()
    assert name.startswith(f"{TS}_") and name.endswith(".pdf")
    if now.month == 1:
        assert "December" in name
    else:
        assert MESIACE[now.month-2] in name
    print(f"  {name}")
    print("  [OK] PASSED")

def test_name_all_months():
    """Vsetkych 12 mesiacov."""
    print("\n[TEST 8] Nazov - 12 mesiacov")
    for i, m in enumerate(MESIACE):
        n = generate_output_filename("KM", i, 2025)
        assert n == f"KM_Výplatná Páska_{m}_2025.pdf"
    print("  [OK] PASSED")

# === KONVERZIA TESTY ===

def test_conv_docx():
    """Konverzia .docx -> zaheslovane PDF."""
    print("\n[TEST 9] Konverzia .docx")
    p = os.path.join(_dir, "t.docx")
    doc = Document()
    doc.add_heading('Test', level=1)
    doc.add_paragraph('Mzda: 1500 EUR')
    doc.save(p)
    r, e = convert_file(p, _dir, TS, 0, 2025)
    assert r and os.path.exists(r), f"Err: {e}"
    reader = PdfReader(r)
    assert reader.is_encrypted
    reader.decrypt(TP)
    assert "1500" in reader.pages[0].extract_text()
    os.remove(p)
    os.remove(r)
    print("  [OK] PASSED")

def test_conv_pdf():
    """Zaheslovanie .pdf."""
    print("\n[TEST 10] Zaheslovanie .pdf")
    dp = os.path.join(_dir, "t2.docx")
    pp = os.path.join(_dir, "t2.pdf")
    doc = Document()
    doc.add_paragraph("Test PDF content")
    doc.save(dp)
    docx_to_pdf(dp, pp)
    os.remove(dp)
    r, e = convert_file(pp, _dir, TS, 5, 2026)
    assert r and os.path.exists(r), f"Err: {e}"
    reader = PdfReader(r)
    assert reader.is_encrypted
    reader.decrypt(TP)
    assert "Test" in reader.pages[0].extract_text()
    os.remove(pp)
    os.remove(r)
    print("  [OK] PASSED")

# === CHYBOVE SCENARE ===

def test_err_no_password():
    """Heslo nenajdene."""
    print("\n[TEST 11] Chyba - heslo nenajdene")
    r, e = convert_file("fake.docx", _dir, "NOBODY", 0, 2025)
    assert r is None and "nenájdené" in e
    print("  [OK] PASSED")

def test_err_bad_format():
    """Nepodporovany format."""
    print("\n[TEST 12] Chyba - zly format")
    p = os.path.join(_dir, "t.txt")
    with open(p, "w") as f:
        f.write("x")
    r, e = convert_file(p, _dir, TS, 0, 2025)
    assert r is None and "Nepodporovaný" in e
    os.remove(p)
    print("  [OK] PASSED")

# === RUNNER ===

if __name__ == "__main__":
    tests = [
        test_env_imports, test_env_fonts, test_env_output_folder, test_env_keyring,
        test_emp_crud, test_emp_multiple,
        test_name_auto, test_name_all_months,
        test_conv_docx, test_conv_pdf,
        test_err_no_password, test_err_bad_format
    ]
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    print("=" * 50)
    print(f"  PDF PASSWORD CONVERTER - TESTY | {now}")
    print("=" * 50)
    setup()
    passed = failed = 0
    for t in tests:
        try:
            t()
            passed += 1
        except Exception as ex:
            print(f"  [FAIL] {ex}")
            failed += 1
    cleanup()
    print("\n" + "=" * 50)
    total = len(tests)
    print(f"  VYSLEDOK: {passed} PASSED / {failed} FAILED / {total} TOTAL")
    status = "ALL PASSED" if failed == 0 else "SOME FAILED"
    print(f"  {status}")
    print("=" * 50)

    # Zapis do test_results.txt (poradovnik)
    results_file = os.path.join(_dir, "test_results.txt")
    run_num = 1
    if os.path.exists(results_file):
        with open(results_file, "r", encoding="utf-8") as rf:
            for line in rf:
                if line.startswith("#"):
                    run_num += 1
    with open(results_file, "a", encoding="utf-8") as rf:
        rf.write(f"#{run_num} | {now} | {passed} PASSED / {failed} FAILED / {total} TOTAL | {status}\n")

    sys.exit(0 if failed == 0 else 1)
